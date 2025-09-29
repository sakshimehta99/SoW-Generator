import pandas as pd
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import tempfile
import os

class SOWGenerator:
    def __init__(self):
        self.replacements = {}
        self.phase_mapping = {}
        self.variable_dict = {}
        self.budget_placeholders = {}
    
    def process_excel_file(self, excel_file):
        """Process the uploaded Excel file and extract variables and budget data"""
        # Reset internal state
        self.replacements = {}
        self.phase_mapping = {}
        self.variable_dict = {}
        self.budget_placeholders = {}
        
        # Save to temporary file for pandas processing
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            tmp_file.write(excel_file.getvalue())
            temp_path = tmp_file.name
        
        try:
            # Load Variables sheet
            variable_df = pd.read_excel(temp_path, sheet_name="Variables", dtype=str)
            
            # Convert into dict (placeholders are already in curly braces in Excel)
            variable_dict = pd.Series(
                variable_df["Example Value"].values,
                index=variable_df["Variable Component"]
            ).to_dict()
            
            # Handle date formatting
            for k, v in variable_dict.items():
                if isinstance(v, (datetime, pd.Timestamp)):
                    variable_dict[k] = v.strftime("%d-%b-%y")
                elif isinstance(v, str) and v not in ['nan', '', None]:
                    # Try to parse string dates
                    try:
                        if ' 00:00:00' in v:
                            parsed_date = datetime.strptime(v, "%Y-%m-%d %H:%M:%S")
                            variable_dict[k] = parsed_date.strftime("%d-%b-%y")
                        elif len(v) == 10 and '-' in v:
                            parsed_date = datetime.strptime(v, "%Y-%m-%d")
                            variable_dict[k] = parsed_date.strftime("%d-%b-%y")
                    except (ValueError, TypeError):
                        pass
            
            # Load Budget sheet
            budget_df = pd.read_excel(temp_path, sheet_name="Budget", header=None)
            
            phase_mapping = {}
            for i in range(len(budget_df)):
                phase = str(budget_df.iloc[i, 0]).strip()
                amount = budget_df.iloc[i, 5]
                
                if not phase or phase.lower() in ["nan", ""] or i == 0:
                    continue
                
                if pd.notna(amount) and str(amount).strip() != "":
                    try:
                        amount_val = float(amount)
                        formatted_amount = f"${int(amount_val):,}"
                        phase_mapping[phase.lower()] = formatted_amount
                    except (ValueError, TypeError):
                        phase_mapping[phase.lower()] = ""
                else:
                    phase_mapping[phase.lower()] = ""
            
            # Map Excel phases to placeholders dynamically
            budget_placeholders = {}
            for phase, amount in phase_mapping.items():
                if phase == "total":
                    budget_placeholders["{TOTAL}"] = amount
                    budget_placeholders["{PROJECTTOTAL}"] = amount
                else:
                    placeholder = "{" + phase.upper().replace(" ", "").replace("&", "") + "}"
                    budget_placeholders[placeholder] = amount
            
            # Store separately for UI display
            self.variable_dict = variable_dict
            self.budget_placeholders = budget_placeholders
            
            # Merge Variables + Budget
            self.replacements = {**variable_dict, **budget_placeholders}
            self.phase_mapping = phase_mapping
            
            return self.replacements
            
        finally:
            # Clean up temporary file
            os.unlink(temp_path)
    
    def replace_text_in_doc(self, doc, replacements):
        """Replace placeholders in all paragraphs and table cells"""
        # Replace in paragraphs
        for p in doc.paragraphs:
            full_text = p.text
            for key, val in replacements.items():
                if key in full_text:
                    full_text = full_text.replace(key, str(val))
            
            if full_text != p.text:
                # Clear existing runs and set new text
                for run in p.runs:
                    run.clear()
                if p.runs:
                    p.runs[0].text = full_text
                else:
                    p.add_run(full_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Replace in cell paragraphs
                    for paragraph in cell.paragraphs:
                        full_text = paragraph.text
                        for key, val in replacements.items():
                            if key in full_text:
                                full_text = full_text.replace(key, str(val))
                        
                        if full_text != paragraph.text:
                            # Clear existing runs and set new text
                            for run in paragraph.runs:
                                run.clear()
                            if paragraph.runs:
                                paragraph.runs[0].text = full_text
                            else:
                                paragraph.add_run(full_text)
    
    def sync_budget_table(self, doc, phase_mapping):
        """Find and update the budget table with Excel data"""
        for table_idx, table in enumerate(doc.tables):
            # Check if this looks like a budget table
            if len(table.rows) < 2 or len(table.rows[0].cells) < 3:
                continue
            
            # Look for headers that indicate this is the budget table
            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip().lower())
            
            # Find column indices
            phase_col = None
            estimate_col = None
            
            for i, header in enumerate(headers):
                if 'phase' in header:
                    phase_col = i
                elif 'estimate' in header or 'cost' in header:
                    estimate_col = i
            
            if phase_col is None or estimate_col is None:
                continue
            
            # Update existing rows and track which phases we've seen
            seen_phases = set()
            rows_to_update = []
            
            for row_idx, row in enumerate(table.rows[1:], 1):
                if row_idx >= len(table.rows):
                    break
                
                try:
                    phase_cell = row.cells[phase_col]
                    estimate_cell = row.cells[estimate_col]
                    
                    phase_text = phase_cell.text.strip()
                    phase_lower = phase_text.lower()
                    
                    # Check if this phase exists in our mapping
                    matched_phase = None
                    for excel_phase in phase_mapping.keys():
                        if excel_phase in phase_lower or phase_lower in excel_phase:
                            matched_phase = excel_phase
                            break
                    
                    if matched_phase:
                        # Update the estimate with center alignment
                        estimate_cell.text = ""
                        estimate_paragraph = estimate_cell.paragraphs[0]
                        estimate_run = estimate_paragraph.add_run(phase_mapping[matched_phase])
                        estimate_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        seen_phases.add(matched_phase)
                    elif phase_lower not in ['total', 'estimated total']:
                        # This is a placeholder row that should be removed
                        rows_to_update.append((row_idx, 'remove'))
                
                except IndexError:
                    continue
            
            # Remove placeholder rows (in reverse order to maintain indices)
            for row_idx, action in reversed(rows_to_update):
                if action == 'remove':
                    try:
                        tbl = table._tbl
                        tr = table.rows[row_idx]._tr
                        tbl.remove(tr)
                    except Exception:
                        pass
            
            # Add missing phases
            total_row_idx = None
            for row_idx, row in enumerate(table.rows[1:], 1):
                try:
                    phase_text = row.cells[phase_col].text.strip().lower()
                    if 'total' in phase_text:
                        total_row_idx = row_idx
                        break
                except IndexError:
                    continue
            
            # Insert missing phases before total row
            for excel_phase, amount in phase_mapping.items():
                if excel_phase not in seen_phases and excel_phase != 'total':
                    try:
                        # Add new row
                        if total_row_idx:
                            # Insert before total
                            new_row = table.add_row()
                            # Move the new row before total
                            tbl = table._tbl
                            new_tr = new_row._tr
                            total_tr = table.rows[total_row_idx]._tr
                            tbl.remove(new_tr)
                            tbl.insert(tbl.index(total_tr), new_tr)
                        else:
                            new_row = table.add_row()
                        
                        # Set the content with proper formatting
                        phase_cell = new_row.cells[phase_col]
                        estimate_cell = new_row.cells[estimate_col]
                        
                        # Clear existing content and add formatted text
                        phase_cell.text = ""
                        phase_paragraph = phase_cell.paragraphs[0]
                        phase_run = phase_paragraph.add_run(excel_phase.title())
                        phase_run.bold = True
                        phase_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Set estimate with proper alignment
                        estimate_cell.text = ""
                        estimate_paragraph = estimate_cell.paragraphs[0]
                        estimate_run = estimate_paragraph.add_run(amount)
                        estimate_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    except Exception:
                        pass
            
            # Update total if it exists
            if 'total' in phase_mapping:
                for row in table.rows[1:]:
                    try:
                        phase_text = row.cells[phase_col].text.strip().lower()
                        if 'total' in phase_text:
                            # Update total with center alignment
                            total_estimate_cell = row.cells[estimate_col]
                            total_estimate_cell.text = ""
                            total_paragraph = total_estimate_cell.paragraphs[0]
                            total_run = total_paragraph.add_run(phase_mapping['total'])
                            total_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            break
                    except IndexError:
                        continue
            
            return  # Only update the first budget table found
    
    def generate_preview(self, excel_path, word_path):
        """Generate a text preview of the document"""
        # Process files
        with open(excel_path, 'rb') as f:
            excel_content = io.BytesIO(f.read())
        
        replacements = self.process_excel_file(excel_content)
        
        # Load and process document
        doc = Document(word_path)
        self.replace_text_in_doc(doc, replacements)
        self.sync_budget_table(doc, self.phase_mapping)
        
        # Extract text for preview
        preview_text = []
        preview_text.append("DOCUMENT PREVIEW")
        preview_text.append("=" * 50)
        preview_text.append("")
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                preview_text.append(paragraph.text)
        
        # Extract table content
        for table_idx, table in enumerate(doc.tables):
            preview_text.append(f"\n--- TABLE {table_idx + 1} ---")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    preview_text.append(row_text)
        
        return "\n".join(preview_text)
    
    def generate_sow(self, excel_path, word_path):
        """Generate the final SOW document and return as bytes buffer"""
        # Process files
        with open(excel_path, 'rb') as f:
            excel_content = io.BytesIO(f.read())
        
        replacements = self.process_excel_file(excel_content)
        
        # Load and process document
        doc = Document(word_path)
        self.replace_text_in_doc(doc, replacements)
        self.sync_budget_table(doc, self.phase_mapping)
        
        # Save to buffer
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        
        return output_buffer
