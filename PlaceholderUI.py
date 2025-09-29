import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import tempfile
import os

# Page configuration
st.set_page_config(
    page_title="SOW Generator",
    page_icon="📄",
    layout="wide"
)

# Custom CSS for better styling
st.markdown("""
    <style>
    .main-header {
        font-size: 50px !important;;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.6rem !important;
        font-weight: bold;
        color: #2c3e50;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if 'replacements' not in st.session_state:
    st.session_state.replacements = None
if 'generated_doc' not in st.session_state:
    st.session_state.generated_doc = None
if 'output_filename' not in st.session_state:
    st.session_state.output_filename = None

def load_variables_from_excel(excel_file):
    """Load variables from Excel file and create replacement dictionary"""
    try:
        # Load Variables sheet
        variable_df = pd.read_excel(excel_file, sheet_name="Variables", dtype=str)
        
        # Convert into dict
        variable_dict = pd.Series(
            variable_df["Example Value"].values,
            index=variable_df["Variable Component"]
        ).to_dict()
        
        # Handle date formatting
        for k, v in variable_dict.items():
            if isinstance(v, (datetime, pd.Timestamp)):
                variable_dict[k] = v.strftime("%d-%b-%y")
            elif isinstance(v, str) and v not in ['nan', '', None]:
                try:
                    if ' 00:00:00' in v:
                        parsed_date = datetime.strptime(v, "%Y-%m-%d %H:%M:%S")
                        variable_dict[k] = parsed_date.strftime("%d-%b-%y")
                    elif len(v) == 10 and '-' in v:
                        parsed_date = datetime.strptime(v, "%Y-%m-%d")
                        variable_dict[k] = parsed_date.strftime("%d-%b-%y")
                except (ValueError, TypeError):
                    pass
        
        # Load Budget sheet
        budget_df = pd.read_excel(excel_file, sheet_name="Budget", header=None)
        
        phase_mapping = {}
        for i in range(len(budget_df)):
            phase = str(budget_df.iloc[i, 0]).strip()
            amount = budget_df.iloc[i, 5]
            
            if not phase or phase.lower() in ["nan", ""] or i == 0:
                continue
            
            if pd.notna(amount) and str(amount).strip() != "":
                try:
                    amount_val = float(amount)
                    formatted_amount = f"${int(amount_val):,}"
                    phase_mapping[phase.lower()] = formatted_amount
                except (ValueError, TypeError):
                    phase_mapping[phase.lower()] = ""
            else:
                phase_mapping[phase.lower()] = ""
        
        # Map Excel phases to placeholders
        budget_placeholders = {}
        for phase, amount in phase_mapping.items():
            if phase == "total":
                budget_placeholders["{TOTAL}"] = amount
                budget_placeholders["{PROJECTTOTAL}"] = amount
            else:
                placeholder = "{" + phase.upper().replace(" ", "").replace("&", "") + "}"
                budget_placeholders[placeholder] = amount
        
        # Merge Variables + Budget
        replacements = {**variable_dict, **budget_placeholders}
        
        return replacements, phase_mapping
        
    except Exception as e:
        st.error(f"Error loading Excel file: {str(e)}")
        return None, None

def replace_text_in_doc(doc, replacements):
    """Replace placeholders in all paragraphs and table cells"""
    for p in doc.paragraphs:
        full_text = p.text
        for key, val in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, str(val))
        
        if full_text != p.text:
            for run in p.runs:
                run.clear()
            if p.runs:
                p.runs[0].text = full_text
            else:
                p.add_run(full_text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    for key, val in replacements.items():
                        if key in full_text:
                            full_text = full_text.replace(key, str(val))
                    
                    if full_text != paragraph.text:
                        for run in paragraph.runs:
                            run.clear()
                        if paragraph.runs:
                            paragraph.runs[0].text = full_text
                        else:
                            paragraph.add_run(full_text)

def sync_budget_table(doc, phase_mapping):
    """Find and update the budget table with Excel data"""
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2 or len(table.rows[0].cells) < 3:
            continue
        
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        
        phase_col = None
        estimate_col = None
        
        for i, header in enumerate(headers):
            if 'phase' in header:
                phase_col = i
            elif 'estimate' in header or 'cost' in header:
                estimate_col = i
        
        if phase_col is None or estimate_col is None:
            continue
        
        seen_phases = set()
        rows_to_update = []
        
        for row_idx, row in enumerate(table.rows[1:], 1):
            if row_idx >= len(table.rows):
                break
            
            try:
                phase_cell = row.cells[phase_col]
                estimate_cell = row.cells[estimate_col]
                
                phase_text = phase_cell.text.strip()
                phase_lower = phase_text.lower()
                
                matched_phase = None
                for excel_phase in phase_mapping.keys():
                    if excel_phase in phase_lower or phase_lower in excel_phase:
                        matched_phase = excel_phase
                        break
                
                if matched_phase:
                    estimate_cell.text = ""
                    estimate_paragraph = estimate_cell.paragraphs[0]
                    estimate_run = estimate_paragraph.add_run(phase_mapping[matched_phase])
                    estimate_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    seen_phases.add(matched_phase)
                elif phase_lower not in ['total', 'estimated total']:
                    rows_to_update.append((row_idx, 'remove'))
            except IndexError:
                continue
        
        for row_idx, action in reversed(rows_to_update):
            if action == 'remove':
                try:
                    tbl = table._tbl
                    tr = table.rows[row_idx]._tr
                    tbl.remove(tr)
                except Exception:
                    pass
        
        total_row_idx = None
        for row_idx, row in enumerate(table.rows[1:], 1):
            try:
                phase_text = row.cells[phase_col].text.strip().lower()
                if 'total' in phase_text:
                    total_row_idx = row_idx
                    break
            except IndexError:
                continue
        
        for excel_phase, amount in phase_mapping.items():
            if excel_phase not in seen_phases and excel_phase != 'total':
                try:
                    if total_row_idx:
                        new_row = table.add_row()
                        tbl = table._tbl
                        new_tr = new_row._tr
                        total_tr = table.rows[total_row_idx]._tr
                        tbl.remove(new_tr)
                        tbl.insert(tbl.index(total_tr), new_tr)
                    else:
                        new_row = table.add_row()
                    
                    phase_cell = new_row.cells[phase_col]
                    estimate_cell = new_row.cells[estimate_col]
                    
                    phase_cell.text = ""
                    phase_paragraph = phase_cell.paragraphs[0]
                    phase_run = phase_paragraph.add_run(excel_phase.title())
                    phase_run.bold = True
                    phase_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    estimate_cell.text = ""
                    estimate_paragraph = estimate_cell.paragraphs[0]
                    estimate_run = estimate_paragraph.add_run(amount)
                    estimate_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
        
        if 'total' in phase_mapping:
            for row in table.rows[1:]:
                try:
                    phase_text = row.cells[phase_col].text.strip().lower()
                    if 'total' in phase_text:
                        total_estimate_cell = row.cells[estimate_col]
                        total_estimate_cell.text = ""
                        total_paragraph = total_estimate_cell.paragraphs[0]
                        total_run = total_paragraph.add_run(phase_mapping['total'])
                        total_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        break
                except IndexError:
                    continue
        
        return

def generate_sow(word_file, replacements, phase_mapping):
    """Generate SOW document"""
    try:
        # Load the Word document from uploaded file
        doc = Document(io.BytesIO(word_file.read()))
        
        # Replace text placeholders
        replace_text_in_doc(doc, replacements)
        
        # Sync budget table
        sync_budget_table(doc, phase_mapping)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
        
    except Exception as e:
        st.error(f"Error generating SOW: {str(e)}")
        return None

# Main UI
st.markdown('<p class="main-header">📄 SOW Generator</p>', unsafe_allow_html=True)
st.markdown("---")

# File Upload Section
col1, col2 = st.columns(2)

with col1:
    st.markdown('<p class="section-header">1️⃣ Upload Excel File</p>', unsafe_allow_html=True)
    excel_file = st.file_uploader(
        "Upload your pricing sheet (Excel)",
        type=['xlsx', 'xls'],
        help="Excel file must contain 'Variables' and 'Budget' sheets"
    )

with col2:
    st.markdown('<p class="section-header">2️⃣ Upload Word Template</p>', unsafe_allow_html=True)
    word_file = st.file_uploader(
        "Upload SOW template (Word)",
        type=['docx'],
        help="Word document with placeholders to be replaced"
    )

# Process Excel file if uploaded
if excel_file:
    replacements, phase_mapping = load_variables_from_excel(excel_file)
    
    if replacements:
        st.session_state.replacements = replacements
        st.session_state.phase_mapping = phase_mapping
        
        st.markdown('<p class="section-header">3️⃣ Placeholder Mappings</p>', unsafe_allow_html=True)
        st.success(f"✅ Loaded {len(replacements)} placeholders from Excel file")
        
        # Display mappings in expandable section
        with st.expander("📋 View All Placeholder Mappings", expanded=False):
            # Create a clean dataframe for display
            display_data = []
            for key, value in replacements.items():
                display_data.append({
                    "Placeholder": key,
                    "Value": value if value not in ['nan', '', None] else "N/A"
                })
            
            df_display = pd.DataFrame(display_data)
            st.dataframe(df_display, use_container_width=True, height=400)
        
        # Search functionality
        search_term = st.text_input("🔍 Search for a specific placeholder", "")
        if search_term:
            filtered = {k: v for k, v in replacements.items() if search_term.lower() in k.lower()}
            if filtered:
                st.write("**Search Results:**")
                for key, value in filtered.items():
                    st.write(f"• **{key}** → {value}")
            else:
                st.warning("No matching placeholders found")

# Generate and Download SOW Section
if excel_file and word_file and st.session_state.replacements:
    st.markdown('<p class="section-header">4️⃣ Generate and Download SOW</p>', unsafe_allow_html=True)
    
    # Generate filename
    client_name = st.session_state.replacements.get('{CLIENT NAME}', 'Client')
    client_name = client_name.replace(' ', '_')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"SOW_{client_name}_{timestamp}.docx"
    
    # Reset word_file pointer before generating
    word_file.seek(0)
    
    # Generate the document
    with st.spinner("Generating your SOW document..."):
        generated_doc = generate_sow(
            word_file,
            st.session_state.replacements,
            st.session_state.phase_mapping
        )
    
    if generated_doc:
        # Single button that downloads immediately
        st.download_button(
            label="🚀 Generate and Download SOW",
            data=generated_doc,
            file_name=output_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
        
        # st.markdown(
        #     '<div class="success-box">✅ SOW is ready! Click the button above to download.</div>',
        #     unsafe_allow_html=True
        # )

# Instructions
with st.expander("ℹ️ How to Use This Tool", expanded=False):
    st.markdown("""
    ### Step-by-Step Guide:
    
    1. **Upload Excel File**: Upload your pricing sheet containing:
       - A `Variables` sheet with placeholder mappings
       - A `Budget` sheet with phase costs
    
    2. **Upload Word Template**: Upload your SOW template document with placeholders (e.g., `{CLIENT NAME}`)
    
    3. **Review Mappings**: Check the placeholder mappings to ensure all values are correct
    
    4. **Generate SOW**: Click the "Generate SOW" button to create your document
    
    5. **Download**: Download the generated SOW document
    
    ### Excel File Requirements:
    - **Variables Sheet**: Must have columns "Variable Component" and "Example Value"
    - **Budget Sheet**: Phase names in column A (0), amounts in column F (5)
    
    ### Date Formatting:
    Dates are automatically formatted as `DD-MMM-YY` (e.g., 13-Aug-25)
    """)

# Footer
st.markdown("---")
st.markdown(
    "<p style='text-align: center; color: #7f8c8d;'>Made with ❤️ for easier SOW generation</p>",
    unsafe_allow_html=True
)